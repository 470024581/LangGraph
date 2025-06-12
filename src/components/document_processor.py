"""
文档处理组件
支持多种文档格式的加载、解析和语义提取
"""
import os
import sqlite3
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional
import asyncio

import pandas as pd
import PyPDF2
from docx import Document
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from loguru import logger

from ..schemas.agent_state import AgentState, DocumentChunk, ProcessingStatus


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self, embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"):
        """初始化文档处理器"""
        self.embedding_model = SentenceTransformer(embedding_model)
        self.vector_store = None
        self.chunk_size = 1000
        self.chunk_overlap = 200
        
    async def process_documents(self, state: AgentState) -> AgentState:
        """处理文档的主入口函数"""
        try:
            state.processing_status = ProcessingStatus.PROCESSING
            state.current_step = "document_processing"
            
            # 加载所有文档
            doc_chunks = await self._load_all_documents()
            state.doc_chunks = doc_chunks
            state.documents_loaded = True
            
            # 构建向量存储
            await self._build_vector_store(doc_chunks)
            state.vector_store_ready = True
            
            # 处理数据库文件
            database_schema = await self._process_database()
            state.database_schema = database_schema
            state.database_loaded = True
            
            state.processing_status = ProcessingStatus.COMPLETED
            state.current_step = "retrieval"
            
            logger.info(f"文档处理完成，共处理 {len(doc_chunks)} 个文档块")
            return state
            
        except Exception as e:
            logger.error(f"文档处理失败: {str(e)}")
            state.processing_status = ProcessingStatus.FAILED
            state.error_message = f"文档处理失败: {str(e)}"
            return state
    
    async def _load_all_documents(self) -> List[DocumentChunk]:
        """加载所有文档"""
        doc_chunks = []
        
        # 文档目录路径
        doc_dir = Path("data/document")
        csv_dir = Path("data/csv")
        
        logger.info(f"开始加载文档，检查目录: {doc_dir.absolute()}")
        
        # 处理document目录
        if doc_dir.exists():
            logger.info(f"找到文档目录: {doc_dir}")
            files_found = list(doc_dir.iterdir())
            logger.info(f"文档目录中发现 {len(files_found)} 个文件/文件夹")
            
            for file_path in files_found:
                logger.info(f"检查文件: {file_path.name}, 是否为文件: {file_path.is_file()}")
                if file_path.is_file():
                    logger.info(f"开始处理文件: {file_path.name} (类型: {file_path.suffix.lower()})")
                    chunks = await self._process_single_file(file_path)
                    logger.info(f"文件 {file_path.name} 处理完成，生成 {len(chunks)} 个文档块")
                    doc_chunks.extend(chunks)
                else:
                    logger.info(f"跳过非文件项: {file_path.name}")
        else:
            logger.warning(f"文档目录不存在: {doc_dir.absolute()}")
        
        # 处理csv目录
        if csv_dir.exists():
            logger.info(f"找到CSV目录: {csv_dir}")
            for file_path in csv_dir.iterdir():
                if file_path.suffix.lower() == '.csv':
                    logger.info(f"开始处理CSV文件: {file_path.name}")
                    chunks = await self._process_csv_file(file_path)
                    logger.info(f"CSV文件 {file_path.name} 处理完成，生成 {len(chunks)} 个文档块")
                    doc_chunks.extend(chunks)
        else:
            logger.warning(f"CSV目录不存在: {csv_dir.absolute()}")
        
        logger.info(f"所有文档加载完成，总共生成 {len(doc_chunks)} 个文档块")
        return doc_chunks
    
    async def _process_single_file(self, file_path: Path) -> List[DocumentChunk]:
        """处理单个文件"""
        file_extension = file_path.suffix.lower()
        
        logger.info(f"准备处理文件: {file_path.name}，文件类型: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                logger.info(f"开始处理PDF文件: {file_path.name}")
                chunks = await self._process_pdf(file_path)
                logger.info(f"PDF文件 {file_path.name} 处理完成，生成 {len(chunks)} 个块")
                return chunks
            elif file_extension == '.docx':
                logger.info(f"开始处理DOCX文件: {file_path.name}")
                chunks = await self._process_docx(file_path)
                logger.info(f"DOCX文件 {file_path.name} 处理完成，生成 {len(chunks)} 个块")
                return chunks
            elif file_extension == '.txt':
                logger.info(f"开始处理TXT文件: {file_path.name}")
                chunks = await self._process_txt(file_path)
                logger.info(f"TXT文件 {file_path.name} 处理完成，生成 {len(chunks)} 个块")
                return chunks
            elif file_extension == '.csv':
                logger.info(f"开始处理CSV文件: {file_path.name}")
                chunks = await self._process_csv_file(file_path)
                logger.info(f"CSV文件 {file_path.name} 处理完成，生成 {len(chunks)} 个块")
                return chunks
            else:
                logger.warning(f"不支持的文件格式: {file_extension} (文件: {file_path.name})")
                return []
                
        except Exception as e:
            logger.error(f"处理文件 {file_path.name} 失败: {str(e)}", exc_info=True)
            return []
    
    async def _process_pdf(self, file_path: Path) -> List[DocumentChunk]:
        """处理PDF文件"""
        chunks = []
        
        logger.info(f"正在读取PDF文件: {file_path.name}")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                logger.info(f"PDF文件 {file_path.name} 共有 {total_pages} 页")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        logger.debug(f"第 {page_num + 1} 页提取到文本，长度: {len(text)} 字符")
                        
                        # 检查是否包含LongLiang相关信息
                        if "longliang" in text.lower() or "long liang" in text.lower():
                            logger.info(f"✅ 第 {page_num + 1} 页包含LongLiang相关信息！")
                            logger.debug(f"   相关内容预览: {text[:500]}...")
                        
                        # 将页面文本分块
                        page_chunks = self._split_text(text)
                        logger.debug(f"第 {page_num + 1} 页分成 {len(page_chunks)} 个块")
                        
                        for chunk_text in page_chunks:
                            chunk = DocumentChunk(
                                content=chunk_text,
                                metadata={
                                    "file_type": "pdf",
                                    "total_pages": total_pages,
                                    "page_number": page_num + 1
                                },
                                chunk_id=str(uuid.uuid4()),
                                source_file=str(file_path),
                                page_number=page_num + 1
                            )
                            chunks.append(chunk)
                            
                            # 检查每个块是否包含LongLiang
                            if "longliang" in chunk_text.lower() or "long liang" in chunk_text.lower():
                                logger.info(f"✅ PDF块包含LongLiang信息 - 页面{page_num + 1}")
                                logger.debug(f"   块内容: {chunk_text[:300]}...")
                    else:
                        logger.debug(f"第 {page_num + 1} 页无有效文本内容")
                
                logger.info(f"PDF文件 {file_path.name} 处理完成，总共生成 {len(chunks)} 个文档块")
        except Exception as e:
            logger.error(f"处理PDF文件 {file_path.name} 时发生错误: {str(e)}", exc_info=True)
            raise
        
        return chunks
    
    async def _process_docx(self, file_path: Path) -> List[DocumentChunk]:
        """处理DOCX文件"""
        chunks = []
        
        logger.info(f"正在读取DOCX文件: {file_path.name}")
        
        try:
            doc = Document(file_path)
            full_text = []
            
            # 提取所有段落文本
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
                    paragraph_count += 1
            
            logger.info(f"DOCX文件 {file_path.name} 提取到 {paragraph_count} 个有效段落")
            
            # 提取表格内容
            table_count = len(doc.tables)
            if table_count > 0:
                logger.info(f"DOCX文件 {file_path.name} 包含 {table_count} 个表格")
                for table_idx, table in enumerate(doc.tables):
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        table_text.append(" | ".join(row_text))
                    full_text.extend(table_text)
                    logger.debug(f"表格 {table_idx + 1} 提取到 {len(table_text)} 行数据")
            
            # 合并所有文本并分块
            document_text = "\n".join(full_text)
            total_text_length = len(document_text)
            logger.info(f"DOCX文件 {file_path.name} 总文本长度: {total_text_length} 字符")
            
            if total_text_length == 0:
                logger.warning(f"DOCX文件 {file_path.name} 没有提取到任何文本内容")
                return chunks
            
            text_chunks = self._split_text(document_text)
            logger.info(f"DOCX文件 {file_path.name} 分成 {len(text_chunks)} 个文本块")
            
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={
                        "file_type": "docx",
                        "paragraph_count": paragraph_count,
                        "table_count": table_count,
                        "chunk_index": i
                    },
                    chunk_id=str(uuid.uuid4()),
                    source_file=str(file_path)
                )
                chunks.append(chunk)
            
            logger.info(f"DOCX文件 {file_path.name} 处理完成，总共生成 {len(chunks)} 个文档块")
        except Exception as e:
            logger.error(f"处理DOCX文件 {file_path.name} 时发生错误: {str(e)}", exc_info=True)
            raise
        
        return chunks
    
    async def _process_txt(self, file_path: Path) -> List[DocumentChunk]:
        """处理TXT文件"""
        chunks = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        text_chunks = self._split_text(text)
        
        for i, chunk_text in enumerate(text_chunks):
            chunk = DocumentChunk(
                content=chunk_text,
                metadata={
                    "file_type": "txt",
                    "file_size": len(text),
                    "chunk_index": i
                },
                chunk_id=str(uuid.uuid4()),
                source_file=str(file_path)
            )
            chunks.append(chunk)
        
        return chunks
    
    async def _process_csv_file(self, file_path: Path) -> List[DocumentChunk]:
        """处理CSV文件，增强了对格式不规范文件的处理能力"""
        chunks = []
        filename = file_path.name
        
        # 步骤 1: 尝试用不同编码打开文件
        encodings_to_try = ['utf-8', 'gbk', 'latin1']
        file_content = None
        detected_encoding = None
        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    file_content = f.read()
                detected_encoding = encoding
                logger.info(f"成功使用 '{encoding}' 编码读取文件: {filename}")
                break
            except UnicodeDecodeError:
                continue

        if not file_content:
            logger.error(f"无法使用 {encodings_to_try} 解码文件: {filename}")
            return [self._create_error_chunk(str(file_path), "无法解码文件")]

        # 使用 io.StringIO 来让 pandas 读取内存中的内容
        from io import StringIO
        file_io = StringIO(file_content)

        try:
            # 步骤 2: 智能解析CSV
            # 使用 'python' 引擎，因为它对不规则文件更宽容
            # on_bad_lines='skip' 会自动跳过格式错误的行
            df = pd.read_csv(file_io, sep=None, engine='python', on_bad_lines='skip')

            # 检查是否成功加载数据
            if df.empty:
                logger.warning(f"文件 {filename} 可能为空或所有行格式均不正确。")
                return [self._create_error_chunk(str(file_path), "文件为空或所有行格式均不正确")]

            # 生成CSV描述信息
            csv_info = f"""
CSV文件信息:
- 文件名: {filename}
- 编码: {detected_encoding}
- 行数: {len(df)}
- 列数: {len(df.columns)}
- 列名: {', '.join(map(str, df.columns.tolist()))}

数据样本:
{df.head().to_string()}

数据类型:
{df.dtypes.to_string()}
"""
            
            logger.info(f"成功处理CSV文件 {filename}，共 {len(df)} 行 {len(df.columns)} 列。已跳过格式错误的行。")
            
            # 创建CSV信息块
            info_chunk = DocumentChunk(
                content=csv_info,
                metadata={
                    "file_type": "csv",
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist()
                },
                chunk_id=str(uuid.uuid4()),
                source_file=str(file_path)
            )
            chunks.append(info_chunk)
            
            # 将每一行转换为文本块
            for idx, row in df.iterrows():
                row_text = f"行 {idx + 1}:\n"
                for col in df.columns:
                    # 确保所有内容都转换为字符串
                    row_text += f"{str(col)}: {str(row[col])}\n"
                
                row_chunk = DocumentChunk(
                    content=row_text,
                    metadata={
                        "file_type": "csv",
                        "row_index": idx,
                        "chunk_type": "data_row"
                    },
                    chunk_id=str(uuid.uuid4()),
                    source_file=str(file_path)
                )
                chunks.append(row_chunk)
                
        except Exception as e:
            logger.error(f"处理CSV文件 {filename} 失败: {e}", exc_info=True)
            chunks.append(self._create_error_chunk(str(file_path), str(e)))
        
        return chunks

    def _create_error_chunk(self, file_path_str: str, error_message: str) -> DocumentChunk:
        """创建一个包含错误信息的文档块"""
        return DocumentChunk(
            content=f"文件 {Path(file_path_str).name} 处理失败: {error_message}\n请检查文件格式是否正确。",
            metadata={
                "file_type": "csv",
                "error": True,
                "error_message": error_message
            },
            chunk_id=str(uuid.uuid4()),
            source_file=file_path_str
        )

    async def _process_database(self) -> Dict[str, Any]:
        """处理数据库文件"""
        db_path = Path("data/database/erp.db")
        database_schema = {}
        
        if not db_path.exists():
            logger.warning("数据库文件不存在")
            return database_schema
        
        try:
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # 获取所有表名
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            
            for table_name, in tables:
                # 获取表结构
                cursor.execute(f"PRAGMA table_info({table_name})")
                columns = cursor.fetchall()
                
                # 获取表数据样本
                cursor.execute(f"SELECT * FROM {table_name} LIMIT 5")
                sample_data = cursor.fetchall()
                
                database_schema[table_name] = {
                    "columns": [{"name": col[1], "type": col[2]} for col in columns],
                    "sample_data": sample_data,
                    "column_names": [col[1] for col in columns]
                }
            
            conn.close()
            logger.info(f"数据库处理完成，共发现 {len(tables)} 个表")
            
        except Exception as e:
            logger.error(f"处理数据库失败: {str(e)}")
        
        return database_schema
    
    def _split_text(self, text: str) -> List[str]:
        """文本分块"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # 尝试在句号处分割
            if end < len(text):
                last_period = text.rfind('.', start, end)
                if last_period > start:
                    end = last_period + 1
            
            chunk = text[start:end]
            chunks.append(chunk.strip())
            
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks
    
    async def _build_vector_store(self, doc_chunks: List[DocumentChunk]) -> None:
        """构建向量存储"""
        if not doc_chunks:
            return
        
        # 统计不同类型的文档块
        type_counts = {}
        pdf_chunks_with_longliang = []
        
        for i, chunk in enumerate(doc_chunks):
            chunk_type = chunk.metadata.get('file_type', 'unknown')
            type_counts[chunk_type] = type_counts.get(chunk_type, 0) + 1
            
            # 检查PDF块是否包含LongLiang
            if chunk_type == 'pdf' and ('longliang' in chunk.content.lower() or 'long liang' in chunk.content.lower()):
                pdf_chunks_with_longliang.append((i, chunk.source_file, chunk.content[:200]))
        
        logger.info(f"向量存储构建统计:")
        for chunk_type, count in type_counts.items():
            logger.info(f"  {chunk_type}: {count} 个块")
        
        if pdf_chunks_with_longliang:
            logger.info(f"✅ 发现 {len(pdf_chunks_with_longliang)} 个包含LongLiang的PDF块:")
            for idx, source, preview in pdf_chunks_with_longliang:
                logger.info(f"  索引 {idx}: {source}")
                logger.debug(f"    内容预览: {preview}...")
        else:
            logger.warning("❌ 未发现包含LongLiang的PDF块")
        
        # 提取所有文本内容
        texts = [chunk.content for chunk in doc_chunks]
        
        # 生成嵌入
        logger.info("开始生成文档嵌入向量...")
        embeddings = self.embedding_model.encode(texts)
        logger.info(f"嵌入向量生成完成，维度: {embeddings.shape}")
        
        # 创建FAISS索引
        dimension = embeddings.shape[1]
        self.vector_store = faiss.IndexFlatIP(dimension)
        
        # 标准化向量
        faiss.normalize_L2(embeddings)
        self.vector_store.add(embeddings.astype('float32'))
        
        logger.info(f"向量存储构建完成，共索引 {len(texts)} 个文档块")
    
    def search_similar_chunks(self, query: str, top_k: int = 5) -> List[tuple]:
        """搜索相似文档块"""
        if not self.vector_store:
            logger.warning("向量存储未初始化")
            return []
        
        logger.info(f"开始向量搜索，查询: '{query}', top_k: {top_k}")
        
        # 生成查询向量
        query_embedding = self.embedding_model.encode([query])
        faiss.normalize_L2(query_embedding)
        logger.debug(f"查询向量生成完成，维度: {query_embedding.shape}")
        
        # 搜索
        scores, indices = self.vector_store.search(query_embedding.astype('float32'), top_k)
        
        results = list(zip(indices[0], scores[0]))
        logger.info(f"向量搜索完成，返回 {len(results)} 个结果")
        
        # 详细分析搜索结果
        for i, (idx, score) in enumerate(results):
            logger.debug(f"搜索结果 {i+1}: 索引={idx}, 相似度={score:.4f}")
        
        return results 